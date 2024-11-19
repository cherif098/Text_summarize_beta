import spacy
from PyPDF2 import PdfFileReader
from docx import Document
from fpdf import FPDF
from io import BytesIO

nlp_en = spacy.load('en_core_web_sm')
nlp_fr = spacy.load('fr_core_news_sm')

def generate_summary(text, language, precision):
    nlp = nlp_en if language == 'en' else nlp_fr
    doc = nlp(text)
    sentences = list(doc.sents)
    num_sentences = len(sentences)
    summary_length = int(num_sentences * precision)
    summary = ' '.join([sent.text for sent in sentences[:summary_length]])
    return summary

def extract_text_from_pdf(file):
    pdf = PdfFileReader(file)
    text = ''
    for page_num in range(pdf.getNumPages()):
        text += pdf.getPage(page_num).extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    text = ''
    for para in doc.paragraphs:
        text += para.text
    return text
def extract_grammatical_info(text, language):
    nlp = nlp_en if language == 'en' else nlp_fr
    doc = nlp(text)
    info = []
    for sent in doc.sents:
        subjects = [token.text for token in sent if token.dep_ in ('nsubj', 'nsubjpass')]
        verbs = [token.text for token in sent if token.pos_ == 'VERB']
        objects = [token.text for token in sent if token.dep_ in ('dobj', 'pobj')]
        info.append((subjects, verbs, objects))
    return info

def save_summary_as_pdf(summary):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, summary)
    return pdf.output(dest='S').encode('latin-1')  # Retourne directement les bytes

def save_summary_as_word(summary):
    doc = Document()
    doc.add_paragraph(summary)
    word_output = BytesIO()
    doc.save(word_output)
    word_bytes = word_output.getvalue()
    word_output.close()
    return word_bytes